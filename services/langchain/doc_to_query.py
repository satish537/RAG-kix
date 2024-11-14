import os
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import TextLoader, Docx2txtLoader, PyPDFLoader, UnstructuredExcelLoader, JSONLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.llms.ollama import Ollama
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from services.langchain.embedding import get_embedding_function
from langchain.schema.document import Document
from docx import Document as DocxDocument
from fastapi import HTTPException, status
from fastapi.responses import JSONResponse
from collections import Counter
from utilservice import *
from services.langchain.singleton import OllamaSingleton

ollama_model = OllamaSingleton.get_instance()


# Template for generating prompts to query the language model
PROMPT_TEMPLATE = """
Answer the question based only on the following context.

Context:
{context}

---

Question:
{question} 

---

Do not add any additional phrases like 'Here is the response'. Strictly adhere to the following format:


{text_formate}

"""


PROMPT_TEMPLATE2 = """
Answer the question based only on the following context.

Context:
{context}

---

Question:
{question} 

---

Do not add any additional phrases like 'Here is the response'.

"""


# Function to load a document, split it into chunks, and run a query against the Chroma vector store

# Path to the Chroma database
DB_PATH = "././chroma/langchain"


# Main function to load and process documents, and add them to the database
def load_database(data_path, filename, recording_id):
    documents = load_documents(data_path, filename)      
    chunks = split_documents(documents)
    print("enod of func")
    return add_to_db(chunks, recording_id, data_path, filename)          


# Function to load documents based on their file extension
def load_documents(data_path, filename):
    # Load the .docx file
    print(data_path)
    print(filename)
    doc = DocxDocument(f"{data_path}/{filename}")

    # Initialize a list to store LangChain Document objects
    documents = []

    # Process paragraphs and accumulate text into Document objects
    paragraph_text = ""
    for para in doc.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            if len(para.text) > 500:
                # Split large paragraphs into smaller chunks
                chunks = [para.text[i:i+500] for i in range(0, len(para.text), 500)]
                for chunk in chunks:
                    documents.append(Document(page_content=chunk.strip(), metadata={"source": filename}))
            else:
                paragraph_text += para.text.strip() + " "
                if len(paragraph_text) > 500:
                    # If accumulated text is large enough, create a Document object
                    documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))
                    paragraph_text = ""
    print("I am here")
    # If there's any remaining text that wasn't added to a Document
    if paragraph_text.strip():
        documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))

    # Process tables and create one Document per table
    for table in doc.tables:
        table_name = table.rows[0].cells[0].text.strip()  # Extract table name from the first cell of the first row
        table_content = []

        # Extract headers
        column_headers = [cell.text.strip() for cell in table.rows[0].cells]  # First row as column headers

        # Iterate over the rows, skipping the first row (headers)
        for row in table.rows[1:]:
            row_header = row.cells[0].text.strip()  # First column as row header

            # Pair each cell in the row with the corresponding column header
            for col_idx, cell in enumerate(row.cells[1:]):
                entry = f"[\"{row_header}\" & \"{column_headers[col_idx + 1]}\" value is \"{cell.text.strip()}\"]"
                table_content.append(entry)

        # Combine the entire table's data into one Document object with table name at the start
        combined_content = f"Table: {table_name}\n" + ", ".join(table_content)
        documents.append(Document(page_content=combined_content, metadata={"source": filename}))

    return documents


# Function to split documents into smaller chunks
def split_documents(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=500,
        length_function=len,
        is_separator_regex=False,
    )
    split_docs = text_splitter.split_documents(documents)
    print("Split docs")
    print(split_docs)
    return split_docs


# Function to add document chunks to the database
def add_to_db(document_chunks: list[Document], project_id, data_path, filename, ):
    db = Chroma(
        persist_directory=f"{DB_PATH}/{project_id}", embedding_function=get_embedding_function()
    )
    chunks_with_ids = compute_ids(document_chunks, project_id, filename)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])
    print("chucks")
    new_chunks = []
    for chunk in chunks_with_ids:
        if chunk.metadata["id"] not in existing_ids:
            new_chunks.append(chunk)

    if len(new_chunks):
        print("inside len chunks")
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
        db.persist()
        delete_document(data_path, filename)
        return db
    else:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="This file has already been uploaded", )
    

# Function to compute unique IDs for document chunks
def compute_ids(chunks, project_id, filename):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}:{project_id}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id
    return chunks





def document_prompts(llm_model: str, data_path: str, filename: str, recording_id: str, prompt_obj: dict):

    print(prompt_obj)

    db = load_database(data_path, filename, recording_id)
    prompts_response = {}

    response_language_list = []
    for category, prompt in prompt_obj.items():

        if type(prompt) is dict:

            # Perform a similarity search with the query
            results = db.similarity_search_with_score(prompt['prompt'], k=5)

            # Generate context text from the search results
            context_text = "\n\n---\n\n".join([doc.page_content for doc, _score in results])

            prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
            final_prompt = prompt_template.format(context=context_text, question=prompt['prompt'], text_formate=prompt['textContent'])

            # Initialize the language model and generate a response
            response_text = ollama_model.invoke(final_prompt)
            language_code, language_name = language_detaction(response_text)
            response_language_list.append(language_code)

            prompts_response[category] = {
                'data': response_text,
                'lanCode': language_code
            }

        elif type(prompt) is str:

            # Perform a similarity search with the query
            results = db.similarity_search_with_score(prompt, k=5)

            # Generate context text from the search results
            context_text = "\n\n---\n\n".join([doc.page_content for doc, _score in results])

            prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE2)
            final_prompt = prompt_template.format(context=context_text, question=prompt)

            # Initialize the language model and generate a response
            response_text = ollama_model.invoke(final_prompt)
            language_code, language_name = language_detaction(response_text)
            response_language_list.append(language_code)

            prompts_response[category] = {
                'data': response_text,
                'lanCode': language_code
            }



    # Delete the uploaded document
    delete_document(data_path, filename)
    print(f"{prompts_response} \n\n for {recording_id}")
    
    # Return the response as a JSON response
    return JSONResponse(content=prompts_response, status_code=status.HTTP_200_OK)


    
