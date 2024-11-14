import os
from langchain_community.document_loaders import TextLoader, Docx2txtLoader, PyPDFLoader, UnstructuredExcelLoader, JSONLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.llms.ollama import Ollama
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from services.langchain.embedding import get_embedding_function
from services.langchain.loadDocuments.storeVectors import load_database
from langchain.schema.document import Document
from docx import Document as DocxDocument
from fastapi import HTTPException, status
from fastapi.responses import JSONResponse
from collections import Counter
from datetime import datetime
from utilservice import *



# Main function to load and process documents, and add them to the database
def userUploadedDocument(data_path, projectID, metadata):
    document_objects = load_documents(data_path, metadata) 
    response = load_database(document_objects, projectID, metadata)    

    return response         


# Function to load documents based on their file extension
def load_documents(data_path, metadata):

    file_extension = os.path.splitext(data_path)[1].lower()
    filename = os.path.basename(data_path)


    if file_extension not in ['.docx']:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid file type")

    doc = DocxDocument(data_path)

    documents = []
 
    paragraph_text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            if len(para.text) > 3000:
                
                chunks = [para.text[i:i+3000] for i in range(0, len(para.text), 3000)]
                for chunk in chunks:
                    documents.append(Document(page_content=chunk.strip(), metadata={"source": filename}))
            else:
                paragraph_text += para.text.strip() + " "
                if len(paragraph_text) > 3000:
                    
                    documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))
                    paragraph_text = ""

    if paragraph_text.strip():
        documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))

    for table in doc.tables:
        table_name = table.rows[0].cells[0].text.strip()
        table_content = []

        column_headers = [cell.text.strip() for cell in table.rows[0].cells]

        for row in table.rows[1:]:
            row_header = row.cells[0].text.strip()

            for col_idx, cell in enumerate(row.cells[1:]):
                entry = f"[\"{row_header}\" & \"{column_headers[col_idx + 1]}\" value is \"{cell.text.strip()}\"]"
                table_content.append(entry)

        combined_content = f"Table: {table_name}\n" + ", ".join(table_content)
        documents.append(Document(page_content=combined_content, metadata={"source": filename}))

    # Extract date and time separately
    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%H:%M:%S")

    for chunk in documents:
        chunk.page_content = chunk.page_content + " METADATA " + f" DATE: '{date_str}', TIME: '{time_str}'" + f"{str(metadata)}"

    return documents


