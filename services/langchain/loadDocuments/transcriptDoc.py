import docx
from langchain.schema import Document as LangchainDocument
from services.langchain.loadDocuments.storeVectors import load_database


async def transcriptToChunks(
    docx_file_path: str, 
    projectID: str, 
    metadata: dict
):
    document_objects = create_document_objects(docx_file_path)
    response = load_database(document_objects, projectID, metadata)
    
    return response



def create_document_objects(docx_file):
    doc = docx.Document(docx_file)
    document_objects = []

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            document_objects.append(LangchainDocument(page_content=paragraph_text))

    return document_objects


