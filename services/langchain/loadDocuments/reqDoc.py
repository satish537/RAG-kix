from docx import Document as DocxDocument
from langchain.schema import Document as LangchainDocument
from services.langchain.loadDocuments.storeVectors import load_database
from services.langchain.dataCleanup.projectData import removeDocuments
from datetime import datetime
from utilservice import *


async def reqDoctoChunks(
    doc_path, 
    projectID, 
    metadata, 
    templateUpdate,
    min_words=100, 
    max_words=300
):

    document_objects = await process_paragraphs(doc_path, projectID, metadata, templateUpdate, min_words, max_words)
    response = load_database(document_objects, projectID, metadata)

    return response


def split_long_paragraph(paragraph_text, max_words=500):
    """Split a long paragraph into two parts of approximately equal size."""
    words = paragraph_text.split()
    mid_index = len(words) // 2

    # Check if the paragraph exceeds the maximum word count
    if len(words) > max_words:
        # Find the nearest space to split into two equal parts
        # Adjust mid_index to find a split point
        while mid_index > 0 and words[mid_index] != '':
            mid_index -= 1

        first_half = ' '.join(words[:mid_index])
        second_half = ' '.join(words[mid_index:])

        return first_half.strip(), second_half.strip()
    else:
        return paragraph_text, None  # Return original paragraph and None

async def process_paragraphs(doc_path, projectID, metadata, templateUpdate, min_words=100, max_words=500):

    # Removing Previous Template Data from Database
    if templateUpdate == "True":
        print("flag is true")
        removeDocuments = await removeDocuments(projectID, metadata["templateID"])


    # Open the DOCX file
    doc = DocxDocument(doc_path)

    documents = []
    current_chunk = ""
    current_word_count = 0

    for para in doc.paragraphs:
        paragraph_text = para.text.strip()

        if not paragraph_text:
            continue  # Skip empty paragraphs

        # Split the long paragraph if it exceeds the maximum word count
        paragraph_text, remaining_paragraph = split_long_paragraph(paragraph_text, max_words)

        # Calculate word count for the current paragraph
        paragraph_word_count = len(paragraph_text.split())

        # Append paragraph to the current chunk if it's too small
        if current_word_count + paragraph_word_count < min_words:
            current_chunk += " " + paragraph_text if current_chunk else paragraph_text
            current_word_count += paragraph_word_count
        else:
            # Add the current chunk as a new Document when it has enough words
            if current_chunk:
                documents.append(LangchainDocument(page_content=current_chunk.strip()))

            # Start a new chunk with the current paragraph
            current_chunk = paragraph_text
            current_word_count = paragraph_word_count

        # If there's a remaining paragraph after splitting, handle it
        if remaining_paragraph:
            # Handle the remaining part separately
            remaining_word_count = len(remaining_paragraph.split())

            if current_word_count + remaining_word_count < min_words:
                current_chunk += " " + remaining_paragraph
                current_word_count += remaining_word_count
            else:
                # Add current chunk as a new Document
                if current_chunk:
                    documents.append(LangchainDocument(page_content=current_chunk.strip()))

                # Start a new chunk with the remaining paragraph
                current_chunk = remaining_paragraph
                current_word_count = remaining_word_count

    # Add the last chunk if it's not empty
    if current_chunk:
        documents.append(LangchainDocument(page_content=current_chunk.strip()))
    
    # Extract date and time separately
    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%H:%M:%S")

    for chunk in documents:
        chunk.page_content = chunk.page_content + "METADATA" + f"DATE: '{date_str}', TIME: '{time_str}'" + f"{str(metadata)}"


    return documents


