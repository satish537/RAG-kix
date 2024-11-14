import os
from langchain_community.document_loaders import TextLoader, Docx2txtLoader, PyPDFLoader, UnstructuredExcelLoader, JSONLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema.document import Document
from services.langchain.embedding import get_embedding_function
from langchain_community.vectorstores import Chroma
from fastapi import HTTPException, status
from fastapi.responses import JSONResponse
from docx import Document as DocxDocument
from langchain.schema import Document
from utilservice import *


# Path to the Chroma database
DB_PATH = "././chroma/langchain"


# Main function to load and process documents, and add them to the database
def load_database(data_path, filename, project_id, documentName, versionID):
    documents = load_documents(data_path, filename)
    return add_to_db(documents, project_id, data_path, filename, documentName, versionID)


# Function to load documents based on their file extension
def load_documents(data_path, filename):
    # Load the .docx file
    doc = DocxDocument(f"{data_path}/{filename}")

    # Initialize a list to store LangChain Document objects
    documents = []

    # Process paragraphs and accumulate text into Document objects
    paragraph_text = ""
    for para in doc.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            if len(para.text) > 500:
                # Split large paragraphs into smaller chunks
                chunks = [para.text[i:i+500] for i in range(0, len(para.text), 500)]
                for chunk in chunks:
                    documents.append(Document(page_content=chunk.strip(), metadata={"source": filename}))
            else:
                paragraph_text += para.text.strip() + " "
                if len(paragraph_text) > 500:
                    # If accumulated text is large enough, create a Document object
                    documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))
                    paragraph_text = ""

    # If there's any remaining text that wasn't added to a Document
    if paragraph_text.strip():
        documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))

    # Iterate through each table in the document
    for table in doc.tables:
        # Get the headers from the first row of the table
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_content = []

        # Iterate through each row in the table starting from the second row
        for row in table.rows[1:]:
            row_data = {}
            for index, cell in enumerate(row.cells):
                row_data[headers[index]] = cell.text.strip()  # Map header to cell text
            
            # Format row text
            row_text = ', '.join([f'{header}: {row_data[header]}' for header in headers])
            table_content.append(row_text)

        # Combine all row texts into a single string for this table
        combined_content = '\n'.join(table_content)
        documents.append(Document(page_content=combined_content, metadata={"source": filename}))


    return documents



# Function to split documents into smaller chunks
def split_documents(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=500,
        length_function=len,
        is_separator_regex=False,
    )
    split_docs = text_splitter.split_documents(documents)

    # for doc in split_docs:
    #     doc.page_content = meta_info + doc.page_content
    return split_docs


# Function to add document chunks to the database
def add_to_db(document_chunks: list[Document], project_id, data_path, filename, documentName, versionID):
    db = Chroma(
        persist_directory=f"{DB_PATH}/{project_id}", embedding_function=get_embedding_function()
    )
    chunks_with_ids = compute_ids(document_chunks, project_id, filename, documentName, versionID)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])

    new_chunks = []
    for chunk in chunks_with_ids:
        if chunk.metadata["id"] not in existing_ids:
            new_chunks.append(chunk)

    if len(new_chunks):
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
        db.persist()
        delete_document(data_path, filename)
        return JSONResponse(content="File Upload Successfully", status_code=status.HTTP_201_CREATED)
    else:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="This file has already been uploaded")
    

# Function to compute unique IDs for document chunks
def compute_ids(chunks, project_id, filename, documentName, versionID):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}:{project_id}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id
        chunk.metadata["documentName"] = documentName
        chunk.metadata["versionID"] = versionID
    return chunks
