import json
import docx
from services.langchain.documentGenration.prompt_gen import prompt_template


def extract_headings_and_data(docx_path):
    """
    Extract headings and data under numbered headings from a Word document.

    Args:
        docx_path (str): Path to the Word document (.docx file)

    Returns:
        A dictionary with the following structure:
        {
            'headings': [
                {'level': 1, 'number': '1', 'data': '...'},
                {'level': 2, 'number': '1.1', 'data': '...'},
                {'level': 3, 'number': '1.1.1', 'data': '...'},
                ...
            ]
        }
    """
    doc = docx.Document(docx_path)
    headings = []

    for idx, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split(' ')[1])
            if level <= 10:  # Extract up to level 10 headings
                number = get_heading_number(level, headings)
                name = para.text  # Extract the heading text
                data = ''
                for next_idx in range(idx + 1, len(doc.paragraphs)):
                    next_para = doc.paragraphs[next_idx]
                    if next_para.style.name.startswith('Heading'):
                        break
                    data += next_para.text + '\n'
                headings.append({'level': level, 'number': number,'name':name, 'data': data.strip()})
    return {'headings': headings}

def get_heading_number(level, headings):
    """
    Generate the heading number based on the level and previous headings.

    Args:
        level (int): The level of the heading (1, 2, 3, etc.)
        headings (list): The list of previously extracted headings

    Returns:
        str: The heading number (e.g. '1', '1.1', '1.1.1', etc.)
    """
    if level == 1:
        return str(len([h for h in headings if h['level'] == 1]) + 1)
    else:
        prev_level = level - 1
        prev_headings = [h for h in headings if h['level'] == prev_level]
        if prev_headings:
            prev_number = prev_headings[-1]['number']
            return '{}.{}'.format(prev_number, len([h for h in headings if h['level'] == level and h['number'].startswith(prev_number)]) + 1)
        else:
            return ''

def remove_data_under_headings(doc_path, headings):
    """
    Remove data under headings and add title and prompt from prompt_resp in the original Word document.

    Args:
        doc_path (str): Path to the Word document
        headings (dict): The extracted headings with their levels, numbers, and data
        prompt_resp (list): The list of prompts with titles and prompts
    """
    print(headings)
    heading_str = "\n".join([f"*{heading['name']}" for heading in headings['headings']])

    #heading_str = "\n".join([f"*{heading}" for heading['name'] in headings])
    print("Heading")
    print(heading_str)
    prompt_resp = prompt_template('mistral',heading_str,'As above given array is topics of document \n\n Give me prompts for above topics for getting details regarding it from given transcript using ai model \n\n\nIn response must be given same formate like : [["topic", prompt], ["topic", prompt], ...]',1)
    doc = docx.Document(doc_path)
    #print(prompt_data.body.decode('utf-8'))
    #prompt_resp = json.loads(prompt_data.body)
    print(prompt_resp)
    #prompt_resp = prompt_data.body.decode('utf-8')
    # Extract title and prompt from prompt_resp
    for idx, para in enumerate(doc.paragraphs):
        # Check if the paragraph is a heading
        if para.style.name.startswith('Heading'):
            # Find the corresponding title and prompt
            for item in prompt_resp:
                title = item['title']
                prompt = item['prompt']
                if para.text == title:

                    # Remove the data under the heading
                    prompt_added = False  # Flag to track if the prompt has been added
                    for next_idx in range(idx + 1, len(doc.paragraphs)):
                        next_para = doc.paragraphs[next_idx]
                        if next_para.style.name.startswith('Heading'):
                            break

                        # Clear the text of the paragraph
                        next_para.clear()  # Clear any content

                        # Reset the style to 'Normal'
                        next_para.style = doc.styles['Normal']

                        # Reset paragraph formatting
                        next_para.paragraph_format.space_before = 0
                        next_para.paragraph_format.space_after = 0
                        next_para.paragraph_format.left_indent = 0
                        next_para.paragraph_format.right_indent = 0
                        next_para.paragraph_format.first_line_indent = 0
                        next_para.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                        # Check for bullet/numbering and remove formatting
                        for run in next_para.runs:
                            run.bold = False
                            run.italic = False
                            run.underline = None  # Removes underlining
                            run.font.color.rgb = None  # Remove text color
                            run.font.size = None  # Reset font size
                            run.font.name = None  # Reset font name
                            run.font.bold = False  # Reset bold formatting
                            run.font.italic = False  # Reset italic formatting
                            run.font.underline = None  # Reset underline formatting

                        # Remove bullet and numbering by setting its properties
                        if next_para._element.xpath('.//w:numPr'):
                            for num_pr in next_para._element.xpath('.//w:numPr'):
                                num_pr.getparent().remove(num_pr)

                        if next_para._element.xpath('.//w:buillt'):
                            for bullet in next_para._element.xpath('.//w:buillt'):
                                bullet.getparent().remove(bullet)

                        # Add the prompt only once
                        if not prompt_added:
                            next_para.text = '{' + title + ' - [' + prompt + ']' + '}'
                            prompt_added = True  # Set the flag to True after adding the prompt

                    # Add the title under the heading
                    para.text = f"{title}"
                    break

    # Save the updated document
    doc.save(doc_path)


def save_headings_to_json(headings, json_path):
    """
    Save the extracted headings to a JSON file.

    Args:
        headings (dict): The extracted headings with their levels, numbers, and data
        json_path (str): Path to the JSON file
    """
    print(headings)
    heading_str = "\n".join([f"*{heading['name']}" for heading in headings['headings']])

    #heading_str = "\n".join([f"*{heading}" for heading['name'] in headings])
    print("Heading")
    print(heading_str)
    #prompt_resp = prompt_template('mistral',heading_str,'As above given array is topics of document \n\n Give me prompts for above topics for getting details regarding it from given transcript using ai model \n\n\nIn response must be given same formate like : [["topic", prompt], ["topic", prompt], ...]')
    with open(json_path, 'w') as file:
        json.dump(headings, file, indent=4)
    return json_path
