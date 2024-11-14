import os
from langchain_community.document_loaders import TextLoader, Docx2txtLoader, PyPDFLoader, UnstructuredExcelLoader, JSONLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema.document import Document
from services.langchain.embedding import get_embedding_function
from langchain_community.vectorstores import Chroma
from fastapi import HTTPException, status
from fastapi.responses import JSONResponse
from docx import Document as DocxDocument
from langchain.schema import Document
from utilservice import *


# Path to the Chroma database
DB_PATH = "././chroma/langchain"


# Main function to load and process documents, and add them to the database
def load_database(data_path, filename, project_id, documentName, versionID):
    documents = load_documents(data_path, filename)      # Load documents from file
    chunks = split_documents(documents)         # Split documents into chunks
    return add_to_db(chunks, project_id, data_path, filename, documentName, versionID)          # Add chunks to the database


# Function to load documents based on their file extension
def load_documents(data_path, filename):
    # Load the .docx file
    doc = DocxDocument(f"{data_path}/{filename}")

    # Initialize a list to store LangChain Document objects
    documents = []

    # Process paragraphs and accumulate text into Document objects
    paragraph_text = ""
    for para in doc.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            if len(para.text) > 500:
                # Split large paragraphs into smaller chunks
                chunks = [para.text[i:i+500] for i in range(0, len(para.text), 500)]
                for chunk in chunks:
                    documents.append(Document(page_content=chunk.strip(), metadata={"source": filename}))
            else:
                paragraph_text += para.text.strip() + " "
                if len(paragraph_text) > 500:
                    # If accumulated text is large enough, create a Document object
                    documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))
                    paragraph_text = ""

    # If there's any remaining text that wasn't added to a Document
    if paragraph_text.strip():
        documents.append(Document(page_content=paragraph_text.strip(), metadata={"source": filename}))

    # Process tables and create one Document per table
    for table in doc.tables:
        table_name = table.rows[0].cells[0].text.strip()  # Extract table name from the first cell of the first row
        table_content = []

        # Extract headers
        column_headers = [cell.text.strip() for cell in table.rows[0].cells]  # First row as column headers

        # Iterate over the rows, skipping the first row (headers)
        for row in table.rows[1:]:
            row_header = row.cells[0].text.strip()  # First column as row header

            # Pair each cell in the row with the corresponding column header
            for col_idx, cell in enumerate(row.cells[1:]):
                entry = f"[\"{row_header}\" & \"{column_headers[col_idx + 1]}\" value is \"{cell.text.strip()}\"]"
                table_content.append(entry)

        # Combine the entire table's data into one Document object with table name at the start
        combined_content = f"Table: {table_name}\n" + ", ".join(table_content)
        documents.append(Document(page_content=combined_content, metadata={"source": filename}))

    return documents



# Function to split documents into smaller chunks
def split_documents(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=500,
        length_function=len,
        is_separator_regex=False,
    )
    split_docs = text_splitter.split_documents(documents)

    # for doc in split_docs:
    #     doc.page_content = meta_info + doc.page_content
    print(split_docs)
    print(len(split_docs))
    return split_docs


# Function to add document chunks to the database
def add_to_db(document_chunks: list[Document], project_id, data_path, filename, documentName, versionID):
    db = Chroma(
        persist_directory=f"{DB_PATH}/{project_id}", embedding_function=get_embedding_function()
    )
    chunks_with_ids = compute_ids(document_chunks, project_id, filename, documentName, versionID)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])

    new_chunks = []
    for chunk in chunks_with_ids:
        if chunk.metadata["id"] not in existing_ids:
            new_chunks.append(chunk)

    if len(new_chunks):
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
        db.persist()
        delete_document(data_path, filename)
        return JSONResponse(content="File Upload Successfully", status_code=status.HTTP_201_CREATED)
    else:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="This file has already been uploaded")
    

# Function to compute unique IDs for document chunks
def compute_ids(chunks, project_id, filename, documentName, versionID):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}:{project_id}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id
        chunk.metadata["documentName"] = documentName
        chunk.metadata["versionID"] = versionID
    return chunks
